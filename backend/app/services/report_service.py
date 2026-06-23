"""Multi-format penetration test report generator — PDF / DOCX / HTML / XLSX
Professional-grade reports with executive summary, CVSS ratings, detailed findings,
attack chain analysis, and remediation recommendations.
"""
import uuid, json, os, io, textwrap
from datetime import datetime
from typing import Optional

from sqlalchemy import create_engine, text as sa_text
from sqlalchemy.orm import Session
from app.config import settings

DB_URL = settings.DATABASE_URL.replace("+asyncpg", "")
REPORT_DIR = "/app/data/reports"
os.makedirs(REPORT_DIR, exist_ok=True)

# ─── helpers ──────────────────────────────────────────────────────────

def _load_task_data(task_id: str) -> Optional[dict]:
    """Load scan task + findings + sessions + credentials from DB."""
    engine = create_engine(DB_URL)
    with Session(engine) as db:
        row = db.execute(sa_text("SELECT * FROM scan_tasks WHERE id=:id"), {"id": task_id}).fetchone()
        if not row:
            return None
        d = dict(row._mapping)
        if isinstance(d.get("result"), str):
            d["result"] = json.loads(d["result"])
        # vulnerabilities
        findings = db.execute(sa_text(
            "SELECT * FROM vulnerabilities WHERE task_id=:tid ORDER BY discovered_at DESC"
        ), {"tid": task_id}).fetchall()
        d["findings"] = [dict(f._mapping) for f in findings]
        # Load sessions from task result JSON
        result_data = d.get("result", {})
        if isinstance(result_data, str):
            try:
                result_data = json.loads(result_data)
            except json.JSONDecodeError:
                result_data = {}
        if isinstance(result_data, dict):
            sessions_raw = result_data.get("sessions", []) or result_data.get("sessions_created", [])
        else:
            sessions_raw = []
        d["sessions"] = sessions_raw if isinstance(sessions_raw, list) else []
        credentials_raw = result_data.get("credentials", []) if isinstance(result_data, dict) else []
        d["credentials"] = credentials_raw if isinstance(credentials_raw, list) else []
        # attack logs
        phases_log = result_data.get("phases_log", {}) if isinstance(result_data, dict) else {}
        logs = []
        if isinstance(phases_log, dict):
            for phase_name, phase_data in phases_log.items():
                if isinstance(phase_data, dict):
                    logs.append({
                        "phase": phase_name,
                        "action": phase_data.get("status", "?"),
                        "target": d.get("target", "?"),
                        "result": phase_data.get("findings", {}).get("summary", ""),
                        "created_at": phase_data.get("completed_at", phase_data.get("started_at", ""))
                    })
        elif isinstance(phases_log, list):
            for entry in phases_log:
                if isinstance(entry, dict):
                    logs.append({
                        "phase": entry.get("name", "?"),
                        "action": entry.get("status", "?"),
                        "target": d.get("target", "?"),
                        "result": json.dumps(entry.get("data", {}), ensure_ascii=False)[:100],
                        "created_at": ""
                    })
        d["attack_logs"] = logs
    return d


def _severity_order(s: str) -> int:
    order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    return order.get((s or "info").lower(), 5)


def _findings_count_by_severity(findings: list) -> dict:
    c = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    for f in findings:
        sev = (f.get("severity") or "info").lower()
        if sev in c:
            c[sev] += 1
    return c


def _findings_by_severity(findings: list, sev: str) -> list:
    return [f for f in findings if (f.get("severity") or "").lower() == sev.lower()]


SEV_CN = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危", "info": "信息"}
SEV_COLORS = {"critical": "#F05050", "high": "#F59E0B", "medium": "#EAB308",
              "low": "#3B82F6", "info": "#94A3B8"}


def _auto_executive_summary(data: dict) -> str:
    """Auto-generate a professional executive summary from findings."""
    findings = data.get("findings", [])
    sessions = data.get("sessions", [])
    creds = data.get("credentials", [])
    sev_count = _findings_count_by_severity(findings)
    target = data.get("target", "?")
    total = sum(sev_count.values())
    c, h, m, l = sev_count["critical"], sev_count["high"], sev_count["medium"], sev_count["low"]
    summary_parts = [f"本次渗透测试目标为 {target}。"]
    if total == 0:
        summary_parts.append("未发现可被利用的安全漏洞，目标整体安全性良好。")
    else:
        summary_parts.append(f"在测试过程中共发现 {total} 个安全漏洞，"
                             f"其中严重 {c} 个、高危 {h} 个、中危 {m} 个、低危 {l} 个。")
        if c + h > 0:
            summary_parts.append(f"高危及以上漏洞共计 {c + h} 个，建议优先修复。")
    if sessions:
        summary_parts.append(f"成功获取 {len(sessions)} 个会话控制权"
                             f"（{', '.join(s.get('type','')+':'+s.get('target','') for s in sessions[:5])}），"
                             f"表明已具备在目标网络内横向移动的能力。")
    if creds:
        summary_parts.append(f"发现 {len(creds)} 组有效凭证信息，涉及"
                             f" {len(set(c.get('username','') for c in creds))} 个不同账户。")
    top = sorted(findings, key=lambda x: _severity_order(x.get("severity", "info")))[:3]
    if top:
        summary_parts.append("重点风险：")
        for i, t in enumerate(top[:3], 1):
            sev = SEV_CN.get((t.get("severity") or "").lower(), t.get("severity", ""))
            summary_parts.append(f"  [{i}] {t.get('title','?')}（{sev}级）")
    summary_parts.append("")
    return "\n".join(summary_parts)


def _cvss_level(score: float) -> str:
    if score >= 9.0: return "Critical"
    if score >= 7.0: return "High"
    if score >= 4.0: return "Medium"
    if score >= 0.1: return "Low"
    return "None"


CVSS_RATINGS = [
    ("严重 (Critical)", "9.0 - 10.0", "攻击者可以完全控制系统，造成严重数据泄露或服务中断"),
    ("高危 (High)", "7.0 - 8.9", "攻击者可以获得系统关键控制权或访问敏感数据"),
    ("中危 (Medium)", "4.0 - 6.9", "攻击者可以获取有限信息或造成部分影响"),
    ("低危 (Low)", "0.1 - 3.9", "攻击者仅能获取极小量非敏感信息"),
    ("信息 (Info)", "0.0", "无直接安全风险，但可作为攻击辅助信息"),
]




def _fmt(dt) -> str:
    """Format a datetime value to string."""
    if isinstance(dt, str):
        return dt[:19]
    if hasattr(dt, "strftime"):
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    return str(dt)

def _pdf_safe(text: str, max_len: int = 0) -> str:
    """Replace unsupported chars for PDF with ASCII equivalents."""
    if not text:
        return ""
    safe = text.encode("ascii", errors="xmlcharrefreplace").decode("ascii")
    if max_len and len(safe) > max_len:
        safe = safe[:max_len]
    return safe


# ══════════════════════════════════════════════════════════════════════
#   PDF (fpdf2)
# ══════════════════════════════════════════════════════════════════════

def _gen_pdf(data: dict, filepath: str):
    try:
        from fpdf import FPDF
    except ImportError:
        raise RuntimeError("fpdf2 is required for PDF generation")

    res = data.get("result") or {}
    summary = res if isinstance(res, dict) else {}
    findings = data.get("findings", [])
    attack_logs = data.get("attack_logs", [])
    sessions = data.get("sessions", [])
    credentials = data.get("credentials", [])
    sev_count = _findings_count_by_severity(findings)

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)

    # Monkey-patch to auto-sanitize non-ASCII text for PDF
    _orig_cell = pdf.cell
    _orig_multi = pdf.multi_cell
    def _safe_cell(w, h=0, txt="", **kw):
        return _orig_cell(w, h, _pdf_safe(str(txt)), **kw)
    def _safe_multi(w, h, txt="", **kw):
        return _orig_multi(w, h, _pdf_safe(str(txt)), **kw)
    pdf.cell = _safe_cell
    pdf.multi_cell = _safe_multi

    # ── Cover ──
    pdf.add_page()
    pdf.set_fill_color(15, 23, 42)
    pdf.rect(0, 0, 210, 110, "F")
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 28)
    pdf.set_y(30)
    pdf.cell(0, 15, "Yunjing Security", align="C")
    pdf.ln(16)
    pdf.set_font("Helvetica", "", 18)
    pdf.cell(0, 10, "Penetration Test Report", align="C")
    pdf.ln(14)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, f"Target: {data.get('target', '?')}", align="C")
    pdf.ln(8)
    pdf.cell(0, 8, f"Date: {datetime.now().strftime('%Y-%m-%d')}", align="C")
    pdf.ln(8)
    pdf.cell(0, 8, f"Report ID: {data.get('id', '?')[:12]}", align="C")
    pdf.set_y(130)
    pdf.set_text_color(100, 100, 100)
    pdf.set_font("Helvetica", "I", 9)
    pdf.cell(0, 6, "CONFIDENTIAL - This report contains sensitive security information.", align="C")
    pdf.ln(6)
    pdf.cell(0, 6, "Distribution is restricted to authorized personnel only.", align="C")

    # ── Executive Summary ──
    pdf.add_page()
    pdf.set_text_color(15, 23, 42)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "1. Executive Summary", new_x="LMARGIN")
    pdf.ln(14)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(55, 65, 81)
    exec_text = _auto_executive_summary(data)
    pdf.multi_cell(0, 6, exec_text)
    pdf.ln(4)

    # Severity summary table
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(15, 23, 42)
    pdf.set_text_color(255, 255, 255)
    cols = ["Severity", "Count", "Percentage"]
    widths = [60, 50, 50]
    for i, (c, w) in enumerate(zip(cols, widths)):
        pdf.cell(w, 8, c, border=1, fill=True, align="C")
    pdf.ln()
    total_f = len(findings) or 1
    for sev, label in [("critical", "Critical"), ("high", "High"),
                       ("medium", "Medium"), ("low", "Low"), ("info", "Info")]:
        cnt = sev_count.get(sev, 0)
        pct = f"{cnt / total_f * 100:.1f}%"
        pdf.set_font("Helvetica", "", 9)
        if cnt > 0:
            pdf.set_fill_color(*{
                "critical": (240, 80, 80), "high": (245, 158, 11),
                "medium": (234, 179, 8), "low": (59, 130, 246), "info": (148, 163, 184)
            }[sev])
            pdf.set_text_color(255, 255, 255)
        else:
            pdf.set_fill_color(249, 250, 251)
            pdf.set_text_color(156, 163, 175)
        pdf.cell(widths[0], 7, label, border=1, fill=True)
        pdf.cell(widths[1], 7, str(cnt), border=1, fill=True, align="C")
        pdf.cell(widths[2], 7, pct, border=1, fill=True, align="C")
        pdf.ln()

    # ── CVSS Rating System ──
    pdf.add_page()
    pdf.set_text_color(15, 23, 42)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "2. CVSS Risk Rating System", new_x="LMARGIN")
    pdf.ln(14)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(55, 65, 81)
    pdf.multi_cell(0, 5, "Vulnerabilities are rated using CVSS v3.1 (Common Vulnerability Scoring System). "
                         "The base score is calculated from: Attack Vector (AV), Attack Complexity (AC), "
                         "Privileges Required (PR), User Interaction (UI), Scope (S), Confidentiality (C), "
                         "Integrity (I), and Availability (A).")
    pdf.ln(6)
    cols = ["Rating", "Score Range", "Description"]
    widths = [40, 30, 90]
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(15, 23, 42)
    pdf.set_text_color(255, 255, 255)
    for c, w in zip(cols, widths):
        pdf.cell(w, 8, c, border=1, fill=True, align="C")
    pdf.ln()
    for rating, score_range, desc in CVSS_RATINGS:
        pdf.set_font("Helvetica", "", 9)
        sev_key = rating.split("(")[1].split(")")[0].lower() if "(" in rating else "info"
        c_cn = {"严重": "critical", "高危": "high", "中危": "medium", "低危": "low"}.get(sev_key, "info")
        score_float = float(score_range.split(" - ")[0]) if " - " in score_range else 0
        if score_float >= 7.0:
            pdf.set_fill_color(252, 165, 165)
        elif score_float >= 4.0:
            pdf.set_fill_color(253, 230, 138)
        else:
            pdf.set_fill_color(219, 234, 254)
        pdf.cell(widths[0], 7, rating, border=1, fill=True)
        pdf.cell(widths[1], 7, score_range, border=1, fill=True, align="C")
        pdf.cell(widths[2], 7, desc, border=1, fill=True)
        pdf.ln()

    # ── Findings ──
    pdf.add_page()
    pdf.set_text_color(15, 23, 42)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "3. Vulnerability Details", new_x="LMARGIN")
    pdf.ln(14)

    for idx, f in enumerate(sorted(findings, key=lambda x: _severity_order(x.get("severity", "info"))), 1):
        sev = (f.get("severity") or "info").lower()
        title = f.get("title", "Unknown")
        target = f.get("target", data.get("target", "?"))
        description = f.get("description") or f.get("detail") or f.get("output", "")
        remediation = f.get("remediation") or f.get("solution") or f.get("fix", "")

        # Risk badge
        sev_label = sev.upper()
        pdf.set_font("Helvetica", "B", 10)
        c = SEV_COLORS.get(sev, "#94A3B8")
        r, g, b = int(c[1:3], 16), int(c[3:5], 16), int(c[5:7], 16)
        pdf.set_fill_color(r, g, b)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(30, 7, f"  {sev_label}", fill=True)
        pdf.set_text_color(15, 23, 42)
        pdf.cell(0, 7, f"  #{idx}  {title[:80]}", new_x="LMARGIN")
        pdf.ln(10)

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(55, 65, 81)
        pdf.cell(0, 5, "Target:", new_x="LMARGIN")
        pdf.ln(5)
        pdf.set_font("Courier", "", 8)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 4, target)
        pdf.ln(2)

        if description:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(55, 65, 81)
            pdf.cell(0, 5, "Description:", new_x="LMARGIN")
            pdf.ln(5)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(55, 65, 81)
            pdf.multi_cell(0, 5, (description[:500] + "...") if len(description) > 500 else description)
            pdf.ln(2)

        if remediation:
            pdf.set_font("Helvetica", "B", 9)
            pdf.set_text_color(55, 65, 81)
            pdf.cell(0, 5, "Remediation:", new_x="LMARGIN")
            pdf.ln(5)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(22, 163, 74)
            pdf.multi_cell(0, 5, (remediation[:300] + "...") if len(remediation) > 300 else remediation)
            pdf.ln(4)

    # ── Sessions (if any) ──
    if sessions:
        pdf.add_page()
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 12, "4. Sessions & Access", new_x="LMARGIN")
        pdf.ln(14)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(55, 65, 81)
        for s in sessions:
            s_type = s.get("type", s.get("session_type", "?"))
            s_target = s.get("target", s.get("host", "?"))
            s_user = s.get("username", s.get("user", "?"))
            s_url = s.get("url", "")
            pdf.set_font("Courier", "", 9)
            pdf.set_fill_color(243, 244, 246)
            pdf.cell(0, 6, f"  [{s_type}] {s_user}@{s_target}" + (f"  ({s_url})" if s_url else ""), fill=True)
            pdf.ln(7)

    # ── Attack Chain ──
    if attack_logs:
        pdf.add_page()
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 12, "5. Attack Chain", new_x="LMARGIN")
        pdf.ln(14)
        pdf.set_font("Courier", "", 8)
        pdf.set_text_color(55, 65, 81)
        for log in attack_logs[:30]:
            phase = log.get("phase", "?")
            action = log.get("action", "?")
            pdf.set_font("Courier", "", 8)
            if "failed" in action.lower() or "error" in action.lower():
                pdf.set_text_color(220, 38, 38)
            elif "completed" in action.lower() or "success" in action.lower():
                pdf.set_text_color(22, 163, 74)
            else:
                pdf.set_text_color(55, 65, 81)
            pdf.cell(0, 5, f"  [{phase}] {action}", new_x="LMARGIN")
            pdf.ln(5)

    # ── Appendix ──
    pdf.add_page()
    pdf.set_text_color(15, 23, 42)
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "6. Appendix", new_x="LMARGIN")
    pdf.ln(14)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(55, 65, 81)
    info_items = [
        ("Report ID", data.get("id", "?")[:12]),
        ("Target", data.get("target", "?")),
        ("Scan Type", data.get("scan_type", "?")),
        ("Status", data.get("status", "?")),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Engine", "Yunjing v2.0 - LLM-Driven Decision Loop"),
    ]
    for k, v in info_items:
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(40, 6, k + ":")
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 6, str(v), new_x="LMARGIN")
        pdf.ln(6)

    pdf.output(filepath)
    return filepath


# ══════════════════════════════════════════════════════════════════════
#   DOCX (python-docx)
# ══════════════════════════════════════════════════════════════════════

def _gen_docx(data: dict, filepath: str):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    res = data.get("result") or {}
    summary = res if isinstance(res, dict) else {}
    findings = data.get("findings", [])
    attack_logs = data.get("attack_logs", [])
    sessions = data.get("sessions", [])
    credentials = data.get("credentials", [])
    sev_count = _findings_count_by_severity(findings)

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(15, 23, 42)
        return h

    def _para(text, bold=False, color=None, size=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if bold: run.bold = True
        if color: run.font.color.rgb = color
        if size: run.font.size = Pt(size)
        return p

    # ── Cover ──
    for _ in range(6):
        doc.add_paragraph()
    _para("YUNJING SECURITY", bold=True, size=28)
    _para("Penetration Test Report", size=18)
    doc.add_paragraph()
    _para(f"Target: {data.get('target', '?')}", size=11)
    _para(f"Date: {datetime.now().strftime('%Y-%m-%d')}", size=11)
    _para(f"Report ID: {data.get('id', '?')[:12]}", size=11)
    doc.add_paragraph()
    _para("CONFIDENTIAL", bold=True, color=RGBColor(220, 38, 38), size=9)
    doc.add_page_break()

    # ── 1. Executive Summary ──
    _heading("1. Executive Summary", level=1)
    _para(_auto_executive_summary(data), size=10)

    # Severity table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Severity", "Count", "Percentage"]):
        hdr[i].text = h
    total_f = len(findings) or 1
    for sev, label in [("critical", "Critical"), ("high", "High"),
                       ("medium", "Medium"), ("low", "Low"), ("info", "Info")]:
        cnt = sev_count.get(sev, 0)
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(cnt)
        row[2].text = f"{cnt / total_f * 100:.1f}%"

    doc.add_page_break()

    # ── 2. CVSS Rating System ──
    _heading("2. CVSS Risk Rating System", level=1)
    _para("Vulnerabilities are rated using CVSS v3.1 (Common Vulnerability Scoring System). "
          "The base score considers: Attack Vector, Attack Complexity, Privileges Required, "
          "User Interaction, Scope, Confidentiality, Integrity, and Availability.", size=10)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Rating", "Score Range", "Description"]):
        hdr[i].text = h
    for rating, score_range, desc in CVSS_RATINGS:
        row = table.add_row().cells
        row[0].text = rating
        row[1].text = score_range
        row[2].text = desc

    doc.add_page_break()

    # ── 3. Findings Summary ──
    _heading("3. Findings Summary", level=1)
    _para(f"Total vulnerabilities found: {len(findings)}", bold=True, size=11)
    doc.add_paragraph()
    for sev, label in [("critical", "Critical"), ("high", "High"),
                       ("medium", "Medium"), ("low", "Low"), ("info", "Info")]:
        cnt = sev_count.get(sev, 0)
        if cnt > 0:
            _para(f"  {label}: {cnt}", size=10)

    # ── 4. Vulnerability Details ──
    doc.add_page_break()
    _heading("4. Vulnerability Details", level=1)
    for idx, f in enumerate(sorted(findings, key=lambda x: _severity_order(x.get("severity", "info"))), 1):
        sev = (f.get("severity") or "info").lower()
        title = f.get("title", "Unknown")
        description = f.get("description") or f.get("detail") or f.get("output", "")
        remediation = f.get("remediation") or f.get("solution") or f.get("fix", "")

        _heading(f"#{idx} {title[:60]}", level=2)
        _para(f"Severity: {sev.upper()}", bold=True, color=RGBColor(*{
            "critical": (240, 80, 80), "high": (245, 158, 11),
            "medium": (234, 179, 8), "low": (59, 130, 246), "info": (148, 163, 184)
        }[sev]), size=10)
        _para(f"Target: {f.get('target', data.get('target', '?'))}", size=9)
        if description:
            _para("Description:", bold=True, size=10)
            _para((description[:600] + "...") if len(description) > 600 else description, size=10)
        if remediation:
            _para("Remediation:", bold=True, size=10)
            _para((remediation[:400] + "...") if len(remediation) > 400 else remediation,
                  color=RGBColor(22, 163, 74), size=10)
        doc.add_paragraph()

    # ── 5. Sessions & Credentials ──
    if sessions or credentials:
        doc.add_page_break()
        _heading("5. Sessions & Credentials", level=1)
        if sessions:
            _heading("Sessions Created", level=2)
            for s in sessions:
                s_type = s.get("type", s.get("session_type", "?"))
                s_target = s.get("target", s.get("host", "?"))
                s_user = s.get("username", s.get("user", "?"))
                s_url = s.get("url", "")
                line = f"[{s_type}] {s_user}@{s_target}" + (f" ({s_url})" if s_url else "")
                _para(line, size=9)
        if credentials:
            _heading("Credentials Discovered", level=2)
            for c in credentials:
                c_user = c.get("username", c.get("user", "?"))
                c_pass = c.get("password", c.get("pass", c.get("secret", "?")))
                c_type = c.get("type", "?")
                _para(f"  [{c_type}] {c_user}:{c_pass}", size=9)

    # ── 6. Attack Chain ──
    if attack_logs:
        doc.add_page_break()
        _heading("6. Attack Chain", level=1)
        for log in attack_logs[:30]:
            phase = log.get("phase", "?")
            action = log.get("action", "?")
            _para(f"[{phase}] {action}", size=9)

    # ── 7. Appendix ──
    doc.add_page_break()
    _heading("7. Appendix", level=1)
    info_items = [
        ("Report ID", str(data.get("id", "?"))[:12]),
        ("Target", str(data.get("target", "?"))),
        ("Scan Type", str(data.get("scan_type", "?"))),
        ("Status", str(data.get("status", "?"))),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Engine", "Yunjing v2.0"),
    ]
    for k, v in info_items:
        p = doc.add_paragraph()
        run_k = p.add_run(f"{k}: ")
        run_k.bold = True
        p.add_run(v)

    doc.save(filepath)
    return filepath


# ══════════════════════════════════════════════════════════════════════
#   HTML
# ══════════════════════════════════════════════════════════════════════

HTML_CSS = """
<style>
  @page { size: A4; margin: 20mm; }
  * { box-sizing: border-box; }
  body { font-family: 'Segoe UI', 'Microsoft YaHei', sans-serif; color: #374151; background: #F3F4F6; margin: 0; padding: 20px; }
  .report { max-width: 210mm; margin: 0 auto; background: #FFFFFF; box-shadow: 0 4px 24px rgba(0,0,0,0.08); padding: 40px; }
  .cover { text-align: center; padding: 80px 40px; border-bottom: 3px solid #0F172A; margin-bottom: 30px; }
  .cover h1 { font-size: 28px; color: #0F172A; margin-bottom: 8px; }
  .cover h2 { font-size: 18px; color: #64748B; font-weight: 400; margin-bottom: 24px; }
  .cover .meta { color: #94A3B8; font-size: 10px; }
  h2.section { font-size: 18px; color: #0F172A; border-bottom: 2px solid #E2E8F0; padding-bottom: 6px; margin: 30px 0 16px; }
  h3.sub { font-size: 14px; color: #1E293B; margin: 16px 0 8px; }
  .sev-badge { display: inline-block; padding: 2px 10px; font-size: 11px; font-weight: 700; color: #fff; border-radius: 4px; }
  table { width: 100%; border-collapse: collapse; margin: 12px 0; font-size: 12px; }
  th { background: #0F172A; color: #fff; padding: 8px 12px; text-align: left; }
  td { padding: 8px 12px; border-bottom: 1px solid #E5E7EB; }
  tr:hover td { background: #F9FAFB; }
  .finding { border: 1px solid #E5E7EB; border-radius: 8px; padding: 16px; margin: 12px 0; background: #FAFAFA; }
  .finding h4 { margin: 0 0 8px; font-size: 14px; }
  .finding .desc { font-size: 12px; line-height: 1.6; color: #4B5563; }
  .finding .fix { font-size: 12px; color: #059669; border-left: 3px solid #059669; padding-left: 12px; margin: 8px 0; }
  .session-line { font-family: 'Courier New', monospace; font-size: 11px; background: #F3F4F6; padding: 4px 8px; border-radius: 4px; margin: 4px 0; }
  @media print { body { background: #fff; padding: 0; } .report { box-shadow: none; padding: 0; } }
</style>
"""

def _gen_html(data: dict, filepath: str):
    res = data.get("result") or {}
    summary = res if isinstance(res, dict) else {}
    findings = data.get("findings", [])
    attack_logs = data.get("attack_logs", [])
    sessions = data.get("sessions", [])
    credentials = data.get("credentials", [])
    sev_count = _findings_count_by_severity(findings)

    def css_color(sev):
        return {"critical": "#F05050", "high": "#F59E0B", "medium": "#EAB308",
                "low": "#3B82F6", "info": "#94A3B8"}.get(sev.lower(), "#94A3B8")

    rows_html = "".join(
        f"<tr><td style='background:{css_color(f.get('severity','info'))};color:#fff'>{sev}</td>"
        f"<td>{f.get('title','?')[:60]}</td><td>{f.get('target',data.get('target','?'))}</td></tr>"
        for f in sorted(findings, key=lambda x: _severity_order(x.get('severity','info')))
        for sev in [(f.get('severity') or 'info').upper()]
    )
    session_html = ""
    if sessions:
        session_html = "<h2 class='section'>Sessions & Access</h2>" + "".join(
            f"<div class='session-line'>[{s.get('type','?')}] {s.get('username','?')}@{s.get('target','?')}"
            + (f" ({s.get('url','')})" if s.get('url') else "") + "</div>"
            for s in sessions
        )
    findings_detail = ""
    for idx, f in enumerate(sorted(findings, key=lambda x: _severity_order(x.get('severity','info'))), 1):
        sev = (f.get('severity') or 'info').lower()
        desc = f.get('description') or f.get('detail') or f.get('output', '')
        rem = f.get('remediation') or f.get('solution') or f.get('fix', '')
        desc_div = ""
        rem_div = ""
        if desc:
            d = (desc[:600] + "...") if len(desc) > 600 else desc
            desc_div = "<div class='desc'><strong>Description:</strong><br>" + d + "</div>"
        if rem:
            r = (rem[:400] + "...") if len(rem) > 400 else rem
            rem_div = "<div class='fix'><strong>Remediation:</strong><br>" + r + "</div>"
        findings_detail += f"""
        <div class='finding'>
          <h4><span class='sev-badge' style='background:{css_color(sev)}'>{sev.upper()}</span> #{idx} {f.get('title','?')[:70]}</h4>
          <div class='desc'><strong>Target:</strong> {f.get('target',data.get('target','?'))}</div>
          {desc_div}
          {rem_div}
        </div>"""

    html = f"""<!DOCTYPE html><html lang="zh-CN"><head><meta charset="utf-8">
<title>Security Report - {data.get('target','?')}</title>{HTML_CSS}</head><body>
<div class="report">
<div class="cover"><h1>YUNJING SECURITY</h1><h2>Penetration Test Report</h2>
<p class="meta">Target: {data.get('target','?')} | {datetime.now().strftime('%Y-%m-%d')} | {data.get('id','?')[:12]}</p>
<p class="meta" style="color:#DC2626;font-weight:700">CONFIDENTIAL</p></div>

<h2 class="section">1. Executive Summary</h2>
<p style="font-size:12px;line-height:1.8">{_auto_executive_summary(data).replace(chr(10),'<br>')}</p>
<table><tr><th>Severity</th><th>Count</th><th>Percentage</th></tr>
{''.join(f'<tr><td>{label}</td><td>{cnt}</td><td>{cnt/(len(findings) or 1)*100:.1f}%</td></tr>'
  for sev,label in [("critical","Critical"),("high","High"),("medium","Medium"),("low","Low"),("info","Info")]
  for cnt in [sev_count.get(sev,0)])}</table>

<h2 class="section">2. CVSS Risk Rating System</h2>
<p style="font-size:12px">Vulnerabilities are rated using CVSS v3.1. Scores range from 0.0 (None) to 10.0 (Critical), based on Attack Vector, Complexity, Privileges, User Interaction, Scope, and Impact on CIA triad.</p>
<table><tr><th>Rating</th><th>Score Range</th><th>Description</th></tr>
{''.join(f'<tr><td>{r}</td><td>{s}</td><td>{d}</td></tr>' for r,s,d in CVSS_RATINGS)}</table>

<h2 class="section">3. Findings Summary</h2>
<p>Total vulnerabilities found: <strong>{len(findings)}</strong></p>
<table><tr><th>#</th><th>Vulnerability</th><th>Target</th></tr>{rows_html}</table>

<h2 class="section">4. Vulnerability Details</h2>{findings_detail}

{session_html}

<h2 class="section">5. Attack Chain</h2>
<table><tr><th>Phase</th><th>Action</th><th>Target</th><th>Result</th><th>Time</th></tr>
{''.join(f'<tr><td>{l.get("phase","?")}</td><td>{l.get("action","?")}</td><td>{l.get("target","?")}</td><td>{str(l.get("result",""))[:40]}</td><td>{_fmt(l.get("created_at",""))}</td></tr>' for l in attack_logs[:30])}</table>

<h2 class="section">6. Appendix</h2>
<table><tr><th>Field</th><th>Value</th></tr>
<tr><td>Report ID</td><td>{str(data.get('id','?'))[:12]}</td></tr>
<tr><td>Target</td><td>{str(data.get('target','?'))}</td></tr>
<tr><td>Scan Type</td><td>{str(data.get('scan_type','?'))}</td></tr>
<tr><td>Status</td><td>{str(data.get('status','?'))}</td></tr>
<tr><td>Generated</td><td>{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</td></tr>
<tr><td>Engine</td><td>Yunjing v2.0</td></tr></table>
</div></body></html>"""
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html)
    return filepath


# ══════════════════════════════════════════════════════════════════════
#   XLSX (openpyxl)
# ══════════════════════════════════════════════════════════════════════

def _gen_xlsx(data: dict, filepath: str):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    res = data.get("result") or {}
    findings = data.get("findings", [])
    attack_logs = data.get("attack_logs", [])
    sessions = data.get("sessions", [])
    credentials = data.get("credentials", [])
    sev_count = _findings_count_by_severity(findings)

    wb = Workbook()
    header_font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="0F172A", end_color="0F172A", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"))

    def _write_header(ws, headers):
        for i, h in enumerate(headers, 1):
            cell = ws.cell(row=1, column=i, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border

    # Sheet 1: Overview
    ws1 = wb.active
    ws1.title = "Overview"
    _write_header(ws1, ["Metric", "Value"])
    metrics = [
        ("Target", data.get("target", "?")),
        ("Scan Type", data.get("scan_type", "?")),
        ("Status", data.get("status", "?")),
        ("Total Vulnerabilities", len(findings)),
        ("Critical", sev_count.get("critical", 0)),
        ("High", sev_count.get("high", 0)),
        ("Medium", sev_count.get("medium", 0)),
        ("Low", sev_count.get("low", 0)),
        ("Ports Found", res.get("ports_found", 0) if isinstance(res, dict) else 0),
        ("Sessions Created", len(sessions)),
        ("Credentials Found", len(credentials)),
    ]
    for i, (k, v) in enumerate(metrics, 2):
        ws1.cell(row=i, column=1, value=k).border = thin_border
        c = ws1.cell(row=i, column=2, value=v)
        c.border = thin_border
        c.alignment = Alignment(horizontal="center")
    ws1.column_dimensions["A"].width = 25
    ws1.column_dimensions["B"].width = 20

    # Sheet 2: Vulnerabilities
    ws2 = wb.create_sheet("Vulnerabilities")
    _write_header(ws2, ["ID", "Severity", "Title", "Target", "Description", "Remediation", "Discovered At"])
    for i, f in enumerate(sorted(findings, key=lambda x: _severity_order(x.get("severity","info"))), 2):
        vals = [i-1, (f.get("severity") or "info").upper(), f.get("title","?")[:60],
                f.get("target",data.get("target","?")),
                (f.get("description") or f.get("detail") or f.get("output",""))[:200],
                (f.get("remediation") or f.get("solution") or f.get("fix",""))[:200],
                _fmt(f.get("discovered_at",""))]
        for j, v in enumerate(vals, 1):
            ws2.cell(row=i, column=j, value=v).border = thin_border
    ws2.column_dimensions["A"].width = 6
    ws2.column_dimensions["B"].width = 12
    ws2.column_dimensions["C"].width = 40
    ws2.column_dimensions["D"].width = 20
    ws2.column_dimensions["E"].width = 40
    ws2.column_dimensions["F"].width = 40
    ws2.column_dimensions["G"].width = 20

    # Sheet 3: Attack Chain
    ws3 = wb.create_sheet("Attack Chain")
    _write_header(ws3, ["Phase", "Action", "Target", "Result", "Time"])
    for i, log in enumerate(attack_logs[:50], 2):
        vals = [log.get("phase","?"), log.get("action","?"), log.get("target","?"),
                str(log.get("result",""))[:60], _fmt(log.get("created_at",""))]
        for j, v in enumerate(vals, 1):
            ws3.cell(row=i, column=j, value=v).border = thin_border

    # Sheet 4: Sessions
    ws4 = wb.create_sheet("Sessions")
    _write_header(ws4, ["Type", "Target", "Username", "URL", "Created At"])
    for i, s in enumerate(sessions, 2):
        vals = [s.get("type",s.get("session_type","?")), s.get("target",s.get("host","?")),
                s.get("username",s.get("user","?")), s.get("url",""), _fmt(s.get("created_at",""))]
        for j, v in enumerate(vals, 1):
            ws4.cell(row=i, column=j, value=v).border = thin_border

    # Sheet 5: Credentials
    ws5 = wb.create_sheet("Credentials")
    _write_header(ws5, ["Type", "Username", "Password/Secret", "Target", "Source"])
    for i, c in enumerate(credentials, 2):
        vals = [c.get("type","?"), c.get("username",c.get("user","?")),
                c.get("password",c.get("pass",c.get("secret","?"))),
                c.get("target",data.get("target","?")), c.get("source","?")]
        for j, v in enumerate(vals, 1):
            ws5.cell(row=i, column=j, value=v).border = thin_border

    wb.save(filepath)
    return filepath


# ══════════════════════════════════════════════════════════════════════
#   Main entry point
# ══════════════════════════════════════════════════════════════════════

def generate_report(task_id: str, fmt: str = "pdf") -> dict:
    """Generate a security scan report in the specified format.

    Args:
        task_id: Scan task UUID
        fmt: Output format - "pdf", "docx", "html", or "xlsx"

    Returns:
        dict with "id", "format", "file_path", and "summary"
    """
    import time
    t0 = time.time()
    data = _load_task_data(task_id)
    if not data:
        return {"error": f"Task {task_id} not found"}

    report_id = str(uuid.uuid4())
    ext = fmt.lower()
    filename = f"{report_id}.{ext}"
    filepath = os.path.join(REPORT_DIR, filename)

    generators = {
        "pdf": _gen_pdf,
        "docx": _gen_docx,
        "html": _gen_html,
        "xlsx": _gen_xlsx,
    }
    gen = generators.get(ext)
    if not gen:
        return {"error": f"Unsupported format: {fmt}"}

    gen(data, filepath)
    elapsed = time.time() - t0
    sev_count = _findings_count_by_severity(data.get("findings", []))
    summary = {
        "total": len(data.get("findings", [])),
        "critical": sev_count.get("critical", 0),
        "high": sev_count.get("high", 0),
        "medium": sev_count.get("medium", 0),
        "low": sev_count.get("low", 0),
        "info": sev_count.get("info", 0),
        "ports_found": data.get("result", {}).get("ports_found", 0) if isinstance(data.get("result"), dict) else 0,
        "status": data.get("status", "?"),
        "target": data.get("target", "?"),
        "scan_type": data.get("scan_type", "?"),
        "sessions": len(data.get("sessions", [])),
        "credentials": len(data.get("credentials", [])),
    }

    # Save to reports table
    engine = create_engine(DB_URL)
    with Session(engine) as db:
        db.execute(sa_text("""
            INSERT INTO reports (id, task_id, format, file_path, summary, created_at)
            VALUES (:id, :task_id, :format, :file_path, :summary, :created_at)
            ON CONFLICT (id) DO UPDATE SET
              format=EXCLUDED.format, file_path=EXCLUDED.file_path,
              summary=EXCLUDED.summary
        """), {
            "id": report_id,
            "task_id": task_id,
            "format": ext,
            "file_path": filepath,
            "summary": json.dumps(summary, ensure_ascii=False),
            "created_at": datetime.utcnow(),
        })
        db.commit()

    return {
        "id": report_id,
        "format": ext,
        "file_path": filepath,
        "summary": summary,
        "elapsed": round(elapsed, 2),
    }