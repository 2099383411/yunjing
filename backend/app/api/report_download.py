"""
多格式报告生成器 — PDF/DOCX/HTML/XLSX
"""
import json, logging, io
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from app.database import AsyncSessionLocal
from app.models.user import User
from app.api.deps import optional_user

logger = logging.getLogger(__name__)
router = APIRouter()

@router.get("/reports/{task_id}/download")
async def download_report(task_id: str, format: str = "html", user: User = Depends(optional_user)):
    """下载多格式渗透测试报告"""
    async with AsyncSessionLocal() as sess:
        from app.models.task import ScanTask
        task = await sess.get(ScanTask, task_id)
        if not task:
            raise HTTPException(404, "任务不存在")
        result = task.result or {}
        findings = result.get("findings", []) if isinstance(result, dict) else []
        ports = result.get("ports", []) if isinstance(result, dict) else []
        summary = result.get("summary", {}) if isinstance(result, dict) else {}

        if format == "html":
            content = _build_html_report(task, findings, ports)
            return StreamingResponse(io.BytesIO(content.encode()), media_type="text/html",
                headers={"Content-Disposition": "attachment; filename=report_" + task_id[:8] + ".html"})
        elif format == "xlsx":
            content = _build_xlsx_report(task, findings, ports)
            return StreamingResponse(content, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": "attachment; filename=report_" + task_id[:8] + ".xlsx"})
        elif format == "docx":
            content = _build_docx_report(task, findings, ports)
            return StreamingResponse(content, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=report_" + task_id[:8] + ".docx"})
        elif format == "pdf":
            content = _build_pdf_report(task, findings, ports)
            return StreamingResponse(content, media_type="application/pdf",
                headers={"Content-Disposition": "attachment; filename=report_" + task_id[:8] + ".pdf"})
        else:
            raise HTTPException(400, "Unsupported format: " + format)

def _severity_color(sev):
    return {"critical":"#dc2626","high":"#ea580c","medium":"#ca8a04","low":"#16a34a","info":"#64748b"}.get(sev, "#94a3b8")

def _build_html_report(task, findings, ports):
    vuln_rows = ""
    for f_item in findings:
        if isinstance(f_item, dict):
            vuln_rows += '<tr><td style="color:' + _severity_color(f_item.get("severity","")) + '"><b>' + str(f_item.get("severity","?")) + '</b></td><td>' + str(f_item.get("name","?")) + '</td><td>' + str(f_item.get("description",""))[:200] + '</td></tr>'
    return """<!DOCTYPE html><html><head><meta charset="utf-8"><title>云镜渗透测试报告</title>
<style>body{font-family:-apple-system,sans-serif;max-width:900px;margin:40px auto;padding:20px;color:#1e293b}
h1{color:#1a3c6e;border-bottom:3px solid #0284c7;padding-bottom:8px}
h2{color:#2a5c9e;margin-top:24px}table{width:100%;border-collapse:collapse;margin:12px 0}
th,td{border:1px solid #e2e8f0;padding:8px 12px;text-align:left;font-size:13px}th{background:#f1f5f9}
.footer{margin-top:40px;padding-top:16px;border-top:1px solid #e2e8f0;color:#94a3b8;font-size:11px}</style></head>
<body><h1>云镜·安全检测助手 — 渗透测试报告</h1>
<p><b>目标:</b> """ + str(task.target) + """ | <b>类型:</b> """ + str(task.scan_type) + """ | <b>状态:</b> """ + str(task.status) + """</p>
<h2>开放端口</h2><p>""" + (", ".join(str(p) for p in ports[:20]) if ports else "无") + """</p>
<h2>漏洞发现 (""" + str(len(findings)) + """个)</h2>
<table><tr><th>严重程度</th><th>名称</th><th>描述</th></tr>""" + vuln_rows + """</table>
<div class="footer">云镜·安全检测助手 | """ + datetime.utcnow().strftime("%Y-%m-%d %H:%M") + """</div></body></html>"""

def _build_xlsx_report(task, findings, ports):
    buf = io.BytesIO()
    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "漏洞清单"
        ws.append(["严重程度", "名称", "类型", "描述", "来源"])
        for f_item in findings:
            if isinstance(f_item, dict):
                ws.append([f_item.get("severity",""), f_item.get("name",""), f_item.get("type",""), str(f_item.get("description",""))[:500], f_item.get("source","")])
        wb.save(buf)
        buf.seek(0)
    except ImportError:
        import csv, codecs
        StreamWriter = codecs.lookup('utf-8')[-1]
        wrapper = StreamWriter(buf)
        writer = csv.writer(wrapper)
        writer.writerow(["severity","name","description"])
        for f_item in findings:
            if isinstance(f_item, dict):
                writer.writerow([f_item.get("severity",""), f_item.get("name",""), str(f_item.get("description",""))[:200]])
    return buf

def _build_docx_report(task, findings, ports):
    buf = io.BytesIO()
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("云镜·渗透测试报告", 0)
        doc.add_paragraph("目标: " + str(task.target) + "\n类型: " + str(task.scan_type) + "\n状态: " + str(task.status))
        doc.add_heading("开放端口", level=1)
        doc.add_paragraph(", ".join(str(p) for p in ports[:20]) if ports else "无")
        doc.add_heading("漏洞发现 (" + str(len(findings)) + "个)", level=1)
        for f_item in findings:
            if isinstance(f_item, dict):
                p = doc.add_paragraph()
                run = p.add_run("[" + str(f_item.get("severity","?")) + "] " + str(f_item.get("name","?")))
                run.bold = True
                doc.add_paragraph(str(f_item.get("description",""))[:300])
        doc.save(buf)
        buf.seek(0)
    except ImportError:
        buf.write(b"DOCX requires python-docx")
        buf.seek(0)
    return buf

def _build_pdf_report(task, findings, ports):
    buf = io.BytesIO()
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, "CloudMirror Report", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 8, "Target: " + str(task.target), ln=True)
        pdf.cell(0, 8, "Status: " + str(task.status), ln=True)
        pdf.ln(4)
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Findings: " + str(len(findings)), ln=True)
        pdf.set_font("Helvetica", "", 9)
        for f_item in findings[:50]:
            if isinstance(f_item, dict):
                pdf.multi_cell(0, 5, "[" + str(f_item.get("severity","?")) + "] " + str(f_item.get("name","?")) + ": " + str(f_item.get("description",""))[:150])
        pdf.output(buf)
        buf.seek(0)
    except ImportError:
        buf.write(b"PDF requires fpdf2")
        buf.seek(0)
    return buf
